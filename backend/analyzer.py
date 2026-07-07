from __future__ import annotations

import html
import json
import re
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Iterable

import openpyxl
import xlrd
from docx import Document
from pypdf import PdfReader

from backend.macos_ocr import recognize_image


AUDIO_SUFFIXES = {".m4a", ".mp3", ".wav", ".aac", ".flac"}
IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png", ".heic", ".tif", ".tiff"}
TEXT_SUFFIXES = {".txt", ".csv", ".md", ".html", ".htm"}
KEYWORDS = (
    "學生",
    "人數",
    "目標",
    "校評",
    "評鑑",
    "分數",
    "續費",
    "續約",
    "教材",
    "器材",
    "gnept",
    "研討",
    "輔導",
)


def compact_text(value: str, limit: int = 50_000) -> str:
    value = value.replace("\x00", " ")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()[:limit]


def extract_html_text(raw: str) -> str:
    raw = re.sub(r"<script\b[^>]*>.*?</script>", " ", raw, flags=re.I | re.S)
    raw = re.sub(r"<style\b[^>]*>.*?</style>", " ", raw, flags=re.I | re.S)
    raw = re.sub(r"</(?:td|th|tr|p|div|br|li)>", "\n", raw, flags=re.I)
    return compact_text(html.unescape(re.sub(r"<[^>]+>", " ", raw)))


def extract_xlsx(path: Path) -> str:
    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sections: list[str] = []
    for sheet in workbook.worksheets:
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value not in (None, "")]
            if values:
                rows.append("\t".join(values))
            if len(rows) >= 500:
                break
        sections.append(f"[工作表] {sheet.title}\n" + "\n".join(rows))
    workbook.close()
    return compact_text("\n\n".join(sections))


def extract_xls(path: Path) -> str:
    try:
        workbook = xlrd.open_workbook(path, on_demand=True)
        sections: list[str] = []
        for sheet in workbook.sheets():
            rows: list[str] = []
            for row_index in range(min(sheet.nrows, 500)):
                values = [str(value).strip() for value in sheet.row_values(row_index) if str(value).strip()]
                if values:
                    rows.append("\t".join(values))
            sections.append(f"[工作表] {sheet.name}\n" + "\n".join(rows))
        workbook.release_resources()
        return compact_text("\n\n".join(sections))
    except (xlrd.XLRDError, UnicodeDecodeError):
        raw = path.read_bytes()
        decoded = raw.decode("utf-8", errors="ignore")
        return extract_html_text(decoded)


def extract_docx(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append("\t".join(values))
    return compact_text("\n".join(parts))


def extract_pdf(path: Path) -> str:
    reader = PdfReader(path)
    pages = [(page.extract_text() or "") for page in reader.pages[:100]]
    return compact_text("\n\n".join(pages))


def extract_image(path: Path, _: Path) -> str:
    return compact_text(recognize_image(path))


def extract_file(path: Path, ocr_script: Path) -> dict[str, object]:
    suffix = path.suffix.casefold()
    result: dict[str, object] = {
        "path": str(path),
        "filename": path.name,
        "status": "已解析",
        "method": "",
        "text": "",
        "error": "",
    }
    try:
        if not suffix and path.stat().st_size == 16_384:
            result.update(status="待補檔", method="資料夾外殼", error="未收到資料夾內的實際檔案")
        elif suffix in AUDIO_SUFFIXES:
            result.update(status="待轉錄", method="錄音", error="錄音轉錄模組尚未啟用")
        elif suffix in IMAGE_SUFFIXES:
            result.update(method="macOS Vision OCR", text=extract_image(path, ocr_script))
        elif suffix in {".xlsx", ".xltx", ".xlsm"}:
            result.update(method="Excel", text=extract_xlsx(path))
        elif suffix == ".xls":
            result.update(method="Excel", text=extract_xls(path))
        elif suffix == ".docx":
            result.update(method="Word", text=extract_docx(path))
        elif suffix == ".pdf":
            text = extract_pdf(path)
            if not text:
                result.update(status="待OCR", method="PDF", error="PDF沒有可擷取文字，需逐頁OCR")
            else:
                result.update(method="PDF", text=text)
        elif suffix in TEXT_SUFFIXES:
            raw = path.read_bytes().decode("utf-8", errors="ignore")
            result.update(method="文字", text=extract_html_text(raw) if suffix in {".html", ".htm"} else compact_text(raw))
        else:
            result.update(status="不支援", method="未知格式", error="目前沒有對應解析器")
    except Exception as error:  # analysis failures must not interrupt uploaded originals
        result.update(status="解析失敗", error=str(error)[:500])
    result["characters"] = len(str(result["text"]))
    return result


def evidence_lines(results: Iterable[dict[str, object]]) -> list[dict[str, str]]:
    evidence: list[dict[str, str]] = []
    for result in results:
        filename = str(result["filename"])
        for line in str(result.get("text", "")).splitlines():
            normalized = line.strip()
            lowered = normalized.casefold()
            if len(normalized) < 3 or not any(keyword in lowered for keyword in KEYWORDS):
                continue
            if re.search(r"\d", normalized):
                evidence.append({"file": filename, "text": normalized[:300]})
            if len(evidence) >= 40:
                return evidence
    return evidence


def extract_metrics(results: list[dict[str, object]], school_code: str) -> dict[str, object]:
    text = "\n".join(str(result.get("text", "")) for result in results)
    normalized = text.replace("人数", "人數").replace("学生", "學生").replace("佔", "占")

    def number_after(pattern: str) -> float | None:
        match = re.search(pattern, normalized, flags=re.I)
        return float(match.group(1)) if match else None

    def clean_number(value: float | None) -> int | float | None:
        if value is None:
            return None
        return int(value) if value.is_integer() else value

    trends: list[dict[str, int]] = []
    lines = [line.strip() for line in normalized.splitlines()]
    for year in (2024, 2025, 2026):
        marker_index = next(
            (
                index
                for index, line in enumerate(lines)
                if school_code.casefold() in line.casefold() and f"{year}實際" in line
            ),
            None,
        )
        if marker_index is None:
            continue
        values: list[int] = []
        for line in lines[marker_index + 1 : marker_index + 20]:
            if re.fullmatch(r"\d+(?:\.0)?", line):
                values.append(int(float(line)))
                if len(values) == 12:
                    break
            elif values:
                break
        nonzero = [value for value in values if value > 0]
        if nonzero:
            trends.append({"year": year, "value": round(sum(nonzero) / len(nonzero))})

    gnept_totals: list[int] = []
    for result in results:
        if "gnept" not in str(result["filename"]).casefold():
            continue
        gnept_totals.extend(int(value) for value in re.findall(r"(?:合計|分校小計)\s*(\d+)", str(result.get("text", ""))))

    evaluation_dimensions: list[dict[str, object]] = []
    dimension_labels = (
        ("向度壹", "學制", 35),
        ("向度貳", "品格", 15),
        ("向度參", "e化", 15),
        ("向度肆", "硬規", 15),
        ("向度伍", "市佔率", 20),
    )
    for result in results:
        result_text = str(result.get("text", ""))
        if "校務評鑑成績一覽表" not in result_text:
            continue
        score_pairs = re.findall(
            r"(?<!\d)(\d{1,3}\.\d{2})\s+(\d{1,3}\.\d{2})(?!\d)",
            result_text,
        )
        if len(score_pairs) < 5:
            continue
        candidate_pairs = score_pairs[:5]
        evaluation_dimensions = [
            {
                "key": key,
                "label": label,
                "weight": weight,
                "rawScore": float(raw_score),
                "weightedScore": float(weighted_score),
            }
            for (key, label, weight), (raw_score, weighted_score) in zip(
                dimension_labels, candidate_pairs
            )
        ]
        break

    evaluation_score = clean_number(number_after(r"(?:總分|得分|德分)[：:\s]*([0-9]+(?:\.[0-9]+)?)"))
    if evaluation_dimensions:
        evaluation_score = clean_number(
            round(
                sum(float(item["weightedScore"]) for item in evaluation_dimensions),
                2,
            )
        )

    students = clean_number(number_after(r"上月學生人數[：:\s]*([0-9.]+)"))
    if students is None and trends:
        students = max(trends, key=lambda item: item["year"])["value"]

    return {
        "students": students,
        "averageStudents": clean_number(number_after(r"近1年平均學生人數[：:\s]*([0-9.]+)")),
        "schoolEvaluationScore": evaluation_score,
        "evaluationDimensions": evaluation_dimensions,
        "contractEnd": (re.search(r"合約到期日[：:\s]*([0-9]{4}/[0-9]{2}/[0-9]{2})", normalized) or [None, None])[1],
        "classCount": clean_number(number_after(r"([0-9]+)\s*班")),
        "teacherCards": clean_number(number_after(r"(?:教師|教聯|數聯)卡[：:\s]*([0-9]+)\s*張")),
        "trainingHours": clean_number(number_after(r"近一年培訓時數[：:\s]*([0-9.]+)\s*小時")),
        "marketShare": clean_number(number_after(r"上月市占率[：:\s]*([0-9.]+)%")),
        "mainMaterialRate": clean_number(number_after(r"近1年主幹教材[：:\s]*([0-9.]+)%")),
        "supplementMaterialRate": clean_number(number_after(r"近1年輔助教材[：:\s]*([0-9.]+)%")),
        "gneptTotal": max(gnept_totals) if gnept_totals else None,
        "studentTrend": trends,
    }


def build_analysis(
    school: dict[str, str],
    upload_id: str,
    file_records: list[dict[str, str]],
    root: Path,
    ocr_script: Path,
) -> dict[str, object]:
    results: list[dict[str, object]] = []
    for record in file_records:
        absolute = root / record["stored_path"]
        result = extract_file(absolute, ocr_script)
        result["path"] = record["stored_path"]
        result["category"] = record["predicted_category"]
        result["year"] = record.get("file_year") or "待確認年度"
        results.append(result)

    statuses = Counter(str(result["status"]) for result in results)
    categories = Counter(record["predicted_category"] for record in file_records)
    parsed_count = statuses.get("已解析", 0)
    total_count = len(results)
    completion = round(parsed_count / total_count * 100) if total_count else 0
    missing_categories = [
        category
        for category in ("經營數據", "續約指標數據", "校務評鑑", "輔導紀錄", "續約管理")
        if categories.get(category, 0) == 0
    ]
    pending = [
        {"filename": result["filename"], "status": result["status"], "reason": result["error"]}
        for result in results
        if result["status"] != "已解析"
    ]
    findings = [
        f"本批共 {total_count} 個項目，成功解析 {parsed_count} 個，內容解析完成度 {completion}%。",
        f"已取得的資料類別：{'、'.join(categories.keys()) or '無'}。",
    ]
    if missing_categories:
        findings.append(f"尚缺資料類別：{'、'.join(missing_categories)}。")
    if pending:
        findings.append(f"有 {len(pending)} 個項目需要補檔、OCR、轉錄或人工處理。")
    else:
        findings.append("本批檔案均已完成內容擷取，可進入規範比對與人工覆核。")

    return {
        "schemaVersion": 1,
        "uploadId": upload_id,
        "school": school,
        "generatedAt": datetime.now().astimezone().isoformat(timespec="seconds"),
        "summary": {
            "totalFiles": total_count,
            "parsedFiles": parsed_count,
            "completionPercent": completion,
            "statusCounts": dict(statuses),
            "categoryCounts": dict(categories),
        },
        "findings": findings,
        "managementRequirements": [
            requirement
            for requirement in (
                "補齊未成功上傳的資料夾內容後重新分析。" if statuses.get("待補檔") else "",
                "確認待分類檔案的正確資料類別。" if categories.get("待分類") else "",
                "錄音完成轉錄前，不將口頭承諾列為已完成事項。" if statuses.get("待轉錄") else "",
                "由輔導員覆核擷取數字與原始文件是否一致。",
            )
            if requirement
        ],
        "metrics": extract_metrics(results, school["code"]),
        "evidence": evidence_lines(results),
        "pendingItems": pending,
        "files": results,
    }
